"""
Script to read ISP letter templates
"""
from docx import Document
import sys

def read_docx(file_path):
    """Read DOCX file and extract text"""
    try:
        doc = Document(file_path)
        print(f"\n{'='*80}")
        print(f"FILE: {file_path}")
        print(f"{'='*80}\n")
        
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                print(f"{para.text}")
        
        # Also check tables
        if doc.tables:
            print("\n--- TABLES ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    print(" | ".join(row_text))
                print()
        
    except Exception as e:
        print(f"Error reading {file_path}: {e}")

if __name__ == "__main__":
    # Read Airtel template
    read_docx(r"c:\Users\saheb\Downloads\New FIR\IP Airtel.docx")
    
    # Read Jio template
    read_docx(r"c:\Users\saheb\Downloads\New FIR\supertraders Relaince Jio Letter.docx")
