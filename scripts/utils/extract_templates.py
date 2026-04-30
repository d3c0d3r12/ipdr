"""
Extract ISP letter templates to text files
"""
from docx import Document

def extract_to_txt(docx_path, txt_path):
    """Extract DOCX content to TXT file"""
    try:
        doc = Document(docx_path)
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"{'='*80}\n")
            f.write(f"FILE: {docx_path}\n")
            f.write(f"{'='*80}\n\n")
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(f"{para.text}\n")
            
            # Extract tables
            if doc.tables:
                f.write("\n\n--- TABLES ---\n\n")
                for table_idx, table in enumerate(doc.tables, 1):
                    f.write(f"Table {table_idx}:\n")
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        f.write(" | ".join(row_text) + "\n")
                    f.write("\n")
        
        print(f"✅ Extracted: {txt_path}")
        
    except Exception as e:
        print(f"❌ Error: {e}")

# Extract both templates
extract_to_txt(
    r"c:\Users\saheb\Downloads\New FIR\IP Airtel.docx",
    r"c:\Users\saheb\Downloads\New FIR\airtel_template_extracted.txt"
)

extract_to_txt(
    r"c:\Users\saheb\Downloads\New FIR\supertraders Relaince Jio Letter.docx",
    r"c:\Users\saheb\Downloads\New FIR\jio_template_extracted.txt"
)

print("\n✅ Done! Check the .txt files")
