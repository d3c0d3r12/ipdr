"""
ISP Letter Generator Utility
Generates official letters to ISPs based on pre-defined templates
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os
import zipfile
import pandas as pd
from pathlib import Path
from typing import Dict, List, Optional
import io
from utils.letter_template import substitute

class ISPLetterGenerator:
    """Generate ISP letters with pre-defined templates"""
    
    # ISP Template Mappings
    ISP_TEMPLATES = {
        'airtel': 'airtel',
        'jio': 'jio',
        'reliance': 'jio',
        'reliance jio': 'jio',
        'vi': 'vi',
        'vodafone': 'vi',
        'vodafone idea': 'vi',
        'idea': 'vi',
        'bsnl': 'default',
        'mtnl': 'default',
    }
    
    # Month mapping for Airtel (numeric to 3-letter)
    MONTH_MAP = {
        '01': 'Jan', '02': 'Feb', '03': 'Mar', '04': 'Apr',
        '05': 'May', '06': 'Jun', '07': 'Jul', '08': 'Aug',
        '09': 'Sep', '10': 'Oct', '11': 'Nov', '12': 'Dec'
    }
    
    def __init__(self):
        self.templates_dir = Path(__file__).parent.parent / "templates" / "isp_letters"
    
    def convert_date_to_airtel_format(self, date_str: str) -> str:
        """
        Convert ANY date format to Airtel format: DD-MMM-YYYY
        Handles: DD-MM-YYYY, YYYY-MM-DD, DD:MM:YYYY, DD.MM.YYYY, YYYYMMDD, DD/MM/YYYY
        Output: 28-Jan-2025 (always with 3-letter month)
        """
        try:
            # Remove any existing formatting
            date_clean = date_str.replace(':', '-').replace('/', '-').replace('.', '-').strip()
            
            # Handle YYYYMMDD format (8 digits, no separators)
            if len(date_clean) == 8 and date_clean.isdigit():
                year = date_clean[:4]
                month = date_clean[4:6]
                day = date_clean[6:8]
            # Handle YYYY-MM-DD format
            elif len(date_clean.split('-')[0]) == 4:
                parts = date_clean.split('-')
                year, month, day = parts[0], parts[1], parts[2]
            # Handle DD-MM-YYYY format
            else:
                parts = date_clean.split('-')
                day, month, year = parts[0], parts[1], parts[2]
            
            # Ensure month is zero-padded for lookup
            month = month.zfill(2)
            day = day.zfill(2)
            
            # Convert month to 3-letter format
            month_abbr = self.MONTH_MAP.get(month, month)
            
            return f"{day}-{month_abbr}-{year}"
        except Exception as e:
            # If all parsing fails, return original
            return date_str
    
    def pad_time_to_6_digits(self, time_str: str) -> str:
        """
        Pad time to 6 digits for Jio
        Input: 12532 or 1532 or 532
        Output: 012532 or 001532 or 000532
        """
        try:
            # Remove any colons or spaces
            time_clean = time_str.replace(':', '').replace(' ', '').strip()
            
            # Pad to 6 digits
            time_padded = time_clean.zfill(6)
            
            return time_padded
        except:
            return time_str
        
    def detect_isps_from_zip(self, zip_file_path: str) -> Dict[str, pd.DataFrame]:
        """
        Extract ISP data from ZIP file (from Step 6)
        
        Returns:
            Dict with ISP name as key and DataFrame as value
        """
        isp_data = {}
        
        try:
            with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
                # List all CSV files in ZIP
                csv_files = [f for f in zip_ref.namelist() if f.endswith('.csv')]
                
                for csv_file in csv_files:
                    # Extract ISP name from filename
                    # e.g., "Airtel_Report.csv" -> "Airtel"
                    # e.g., "Reliance_Jio_Report.csv" -> "Reliance Jio"
                    filename = os.path.splitext(os.path.basename(csv_file))[0]
                    
                    # Remove "_Report" suffix if present
                    if filename.endswith('_Report'):
                        isp_name = filename[:-7]  # Remove last 7 characters "_Report"
                    else:
                        isp_name = filename
                    
                    # Replace underscores with spaces for display
                    isp_name = isp_name.replace('_', ' ')
                    
                    # Read CSV data
                    with zip_ref.open(csv_file) as f:
                        df = pd.read_csv(f)
                        
                        # Skip if empty
                        if len(df) == 0:
                            continue
                        
                        isp_data[isp_name] = df
                        
        except Exception as e:
            raise Exception(f"Error reading ZIP file: {str(e)}")
        
        return isp_data
    
    def get_template_type(self, isp_name: str) -> str:
        """
        Get template type for ISP
        
        Returns:
            'airtel', 'jio', 'vi', or 'default'
        """
        isp_lower = isp_name.lower().strip()
        
        # Check for exact match first
        if isp_lower in self.ISP_TEMPLATES:
            return self.ISP_TEMPLATES[isp_lower]
        
        # Check for partial matches (more flexible)
        if 'airtel' in isp_lower or 'bharti' in isp_lower:
            return 'airtel'
        if 'jio' in isp_lower or 'reliance' in isp_lower:
            return 'jio'
        if 'vi' in isp_lower or 'vodafone' in isp_lower or 'idea' in isp_lower:
            return 'vi'
        
        return 'default'
    
    def format_timestamp_for_airtel(self, timestamp_str: str) -> Dict[str, str]:
        """
        Format timestamp for Airtel
        Format: DD-MMM-YYYY HH24:MI:SS (combined in one column)
        Month in 3-letter format (Jan, Feb, Mar, etc.)
        
        Returns:
            {'from_datetime': '28-Jan-2025 14:30:25', 'to_datetime': '28-Jan-2025 15:30:25'}
        """
        try:
            # Parse timestamp
            dt = pd.to_datetime(timestamp_str)
            
            # From datetime with 3-letter month
            from_dt = dt.strftime('%d-%b-%Y %H:%M:%S')
            
            # To datetime (from + 1 hour) with 3-letter month
            to_dt = (dt + timedelta(hours=1)).strftime('%d-%b-%Y %H:%M:%S')
            
            return {
                'from_datetime': from_dt,
                'to_datetime': to_dt
            }
        except:
            return {
                'from_datetime': 'N/A',
                'to_datetime': 'N/A'
            }
    
    def format_timestamp_for_vi(self, timestamp_str: str) -> Dict[str, str]:
        """
        Format timestamp for VI (Vodafone Idea)
        Format: DD:MM:YYYY and HH:MM:SS (IST) (separate columns)
        
        Returns:
            {'from_date': '28:01:2025', 'from_time': '14:30:25', 'to_date': '28:01:2025', 'to_time': '15:30:25'}
        """
        try:
            dt = pd.to_datetime(timestamp_str)
            
            # From date and time
            from_date = dt.strftime('%d:%m:%Y')
            from_time = dt.strftime('%H:%M:%S')
            
            # To date and time (from + 1 hour)
            to_dt = dt + timedelta(hours=1)
            to_date = to_dt.strftime('%d:%m:%Y')
            to_time = to_dt.strftime('%H:%M:%S')
            
            return {
                'from_date': from_date,
                'from_time': from_time,
                'to_date': to_date,
                'to_time': to_time
            }
        except:
            return {
                'from_date': 'N/A',
                'from_time': 'N/A',
                'to_date': 'N/A',
                'to_time': 'N/A'
            }
    
    def format_timestamp_for_jio(self, timestamp_str: str) -> Dict[str, str]:
        """
        Format timestamp for Jio
        Format: YYYYMMDD and HHMMSS (IST) (separate columns)
        Time must be 6 digits (padded with leading zeros)
        
        Returns:
            {'from_date': '20250128', 'from_time': '143025', 'to_date': '20250128', 'to_time': '153025'}
        """
        try:
            dt = pd.to_datetime(timestamp_str)
            
            # From date and time (time padded to 6 digits)
            from_date = dt.strftime('%Y%m%d')
            from_time = dt.strftime('%H%M%S')  # Always 6 digits
            
            # To date and time (from + 1 hour)
            to_dt = dt + timedelta(hours=1)
            to_date = to_dt.strftime('%Y%m%d')
            to_time = to_dt.strftime('%H%M%S')  # Always 6 digits
            
            return {
                'from_date': from_date,
                'from_time': from_time,
                'to_date': to_date,
                'to_time': to_time
            }
        except:
            return {
                'from_date': 'N/A',
                'from_time': 'N/A',
                'to_date': 'N/A',
                'to_time': 'N/A'
            }
    
    def create_jio_txt_file(self, isp_name: str, ip_data: pd.DataFrame, case_details: Dict) -> str:
        """
        Create plain text file for Jio (tab-separated format)
        This is for easy copy-paste into Jio's portal
        Format: Type\tSearch Value\tFrom Date\tFrom Time\tTo Date\tTo Time
        """
        lines = []
        
        # Header line
        lines.append("Type\tSearch Value\tFrom Date (YYYYMMDD)\tFrom Time (HHMMSS IST)\tTo Date (YYYYMMDD)\tTo Time (HHMMSS IST)")
        
        # Data rows
        for idx, row in ip_data.iterrows():
            # Get IP address
            ip_address = str(row.get('Search Value', row.get('ip', '')))
            
            # Get IP type
            ip_type = str(row.get('Type', ''))
            if not ip_type or ip_type == 'nan':
                ip_type = 'IPV6' if ':' in ip_address else 'IPV4'
            
            # Get dates and times
            from_date = str(row.get('From Date', ''))
            from_time = str(row.get('From Time', '')).replace(':', '').zfill(6)
            to_date = str(row.get('To Date', ''))
            to_time = str(row.get('To Time', '')).replace(':', '').zfill(6)
            
            # Add row
            lines.append(f"{ip_type}\t{ip_address}\t{from_date}\t{from_time}\t{to_date}\t{to_time}")
        
        return '\n'.join(lines)
    
    def render_template_to_docx(self, template: dict, isp_name: str,
                                ip_data: pd.DataFrame, case_details: dict) -> Document:
        """Build a letter .docx by walking a template's blocks."""
        values = dict(case_details)
        values["isp_name"] = isp_name

        doc = Document()
        page = template.get("page", {})
        margins = page.get("margins_inches", {})
        default_font = page.get("default_font", "Calibri")
        default_size = page.get("default_size", 10)
        for section in doc.sections:
            section.top_margin = Inches(margins.get("top", 0.5))
            section.bottom_margin = Inches(margins.get("bottom", 0.5))
            section.left_margin = Inches(margins.get("left", 0.75))
            section.right_margin = Inches(margins.get("right", 0.75))

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        for block in template.get("blocks", []):
            btype = block.get("type")
            if btype == "text":
                p = doc.add_paragraph()
                p.alignment = align_map.get(block.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(substitute(block.get("content", ""), values))
                run.bold = bool(block.get("bold"))
                run.italic = bool(block.get("italic"))
                run.font.name = block.get("font") or default_font
                run.font.size = Pt(block.get("size") or default_size)
            elif btype == "list":
                style = block.get("style", "numbered")
                for i, item in enumerate(block.get("items", []), start=1):
                    text = substitute(item, values)
                    prefix = f"{i}. " if style == "numbered" else "• "
                    p = doc.add_paragraph()
                    run = p.add_run(prefix + text)
                    run.font.name = block.get("font") or default_font
                    run.font.size = Pt(block.get("size") or default_size)
            elif btype == "ip_table":
                self._build_ip_table(doc, isp_name, ip_data)
            elif btype == "spacer":
                for _ in range(int(block.get("lines", 1))):
                    doc.add_paragraph()

        return doc

    def _build_ip_table(self, doc: Document, isp_name: str, ip_data: pd.DataFrame) -> None:
        """Append the telco-correct IP table for the ISP to the document."""
        template_type = self.get_template_type(isp_name)

        if template_type == 'airtel':
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text = 'Type', 'Search Value'
            hdr[2].text = 'From Date\n(DD-MMM-YYYY)\n(HH24:MI:SS)'
            hdr[3].text = 'To Date\n(DD-MMM-YYYY)\n(HH24:MI:SS)'
            self._bold_header(hdr, 9)
            for _, row in ip_data.iterrows():
                cells = table.add_row().cells
                ip = str(row.get('Search Value', row.get('ip', '')))
                ip_type = str(row.get('Type', '')) or ''
                if not ip_type or ip_type == 'nan':
                    ip_type = 'IPV6' if ':' in ip else 'IPV4'
                from_date = str(row.get('From Date', ''))
                from_time = str(row.get('From Time', ''))
                if from_date and from_date != 'nan':
                    fda = self.convert_date_to_airtel_format(from_date)
                    tda = self.convert_date_to_airtel_format(str(row.get('To Date', '')))
                    cells[0].text, cells[1].text = ip_type, ip
                    cells[2].text = fda + ' ' + from_time if from_time and from_time != 'nan' else fda
                    to_time = str(row.get('To Time', ''))
                    cells[3].text = tda + ' ' + to_time if to_time and to_time != 'nan' else tda
                else:
                    f = self.format_timestamp_for_airtel(row.get('timestamp', ''))
                    cells[0].text, cells[1].text = ip_type, ip
                    cells[2].text, cells[3].text = f['from_datetime'], f['to_datetime']
                self._set_cell_font(cells, 9)
            return

        # Jio and all others use 6 columns; Jio pads time to 6 digits.
        is_jio = template_type == 'jio'
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = 'Type', 'Search Value'
        if is_jio:
            hdr[2].text, hdr[3].text = 'From Date\nYYYYMMDD', 'From Time\nHHMMSS\n(IST)'
            hdr[4].text, hdr[5].text = 'To Date\nYYYYMMDD', 'To Time\nHHMMSS\n(IST)'
        else:
            hdr[2].text, hdr[3].text = 'From Date\nDD:MM:YYYY', 'From Time\nHH:MM:SS\n(IST)'
            hdr[4].text, hdr[5].text = 'To Date\nDD:MM:YYYY', 'To Time\nHH:MM:SS\n(IST)'
        self._bold_header(hdr, 8)
        for _, row in ip_data.iterrows():
            cells = table.add_row().cells
            ip = str(row.get('Search Value', row.get('ip', '')))
            ip_type = str(row.get('Type', '')) or ''
            if not ip_type or ip_type == 'nan':
                ip_type = 'IPV6' if ':' in ip else 'IPV4'
            from_date = str(row.get('From Date', ''))
            from_time = str(row.get('From Time', ''))
            to_date = str(row.get('To Date', ''))
            to_time = str(row.get('To Time', ''))
            if from_date and from_date != 'nan':
                if is_jio:
                    from_time = self.pad_time_to_6_digits(from_time) if from_time and from_time != 'nan' else '000000'
                    to_time = self.pad_time_to_6_digits(to_time) if to_time and to_time != 'nan' else '000000'
                cells[0].text, cells[1].text = ip_type, ip
                cells[2].text = from_date
                cells[3].text = from_time if from_time and from_time != 'nan' else ''
                cells[4].text = to_date if to_date and to_date != 'nan' else ''
                cells[5].text = to_time if to_time and to_time != 'nan' else ''
            else:
                f = (self.format_timestamp_for_jio if is_jio else self.format_timestamp_for_vi)(row.get('timestamp', ''))
                cells[0].text, cells[1].text = ip_type, ip
                cells[2].text, cells[3].text = f['from_date'], f['from_time']
                cells[4].text, cells[5].text = f['to_date'], f['to_time']
            self._set_cell_font(cells, 8)

    @staticmethod
    def _bold_header(hdr_cells, size_pt: int) -> None:
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(size_pt)

    @staticmethod
    def _set_cell_font(cells, size_pt: int) -> None:
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)

    def generate_letter(self, isp_name: str, ip_data: pd.DataFrame,
                        case_details: dict, template: dict = None) -> Document:
        """Generate a letter for one ISP using the given template (or the default)."""
        from utils.letter_template import DEFAULT_TEMPLATE
        return self.render_template_to_docx(template or DEFAULT_TEMPLATE,
                                            isp_name, ip_data, case_details)

    def generate_all_letters(self, zip_file_path: str, case_details: dict,
                             template: dict = None) -> bytes:
        """Generate all ISP letters from a Step-6 ZIP. Returns ZIP bytes."""
        isp_data = self.detect_isps_from_zip(zip_file_path)
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for isp_name, ip_df in isp_data.items():
                doc = self.generate_letter(isp_name, ip_df, case_details, template)
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                fir = str(case_details.get('fir_number', 'N-A')).replace('/', '-')
                zip_file.writestr(f"{isp_name}_Letter_{fir}.docx", doc_buffer.getvalue())
                if self.get_template_type(isp_name) == 'jio':
                    txt = self.create_jio_txt_file(isp_name, ip_df, case_details)
                    zip_file.writestr(f"{isp_name}_Data_{fir}.txt", txt)
        zip_buffer.seek(0)
        return zip_buffer.getvalue()
