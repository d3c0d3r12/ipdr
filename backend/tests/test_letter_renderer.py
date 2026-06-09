# backend/tests/test_letter_renderer.py
from utils.letter_template import (
    LetterTemplate, PLACEHOLDERS, substitute, DEFAULT_TEMPLATE,
)


def test_substitute_replaces_known_tokens():
    out = substitute("FIR {fir_number} for {isp_name}",
                     {"fir_number": "201/25", "isp_name": "Airtel"})
    assert out == "FIR 201/25 for Airtel"


def test_substitute_missing_value_becomes_empty():
    assert substitute("Hello {officer_name}!", {}) == "Hello !"
    assert substitute("FIR {fir_number}", {"fir_number": "0"}) == "FIR 0"
    assert substitute("FIR {fir_number}", {"fir_number": 2025}) == "FIR 2025"


def test_placeholders_list_has_expected_tokens():
    assert "fir_number" in PLACEHOLDERS
    assert "isp_name" in PLACEHOLDERS
    assert len(PLACEHOLDERS) == 13, "Update this count when adding/removing placeholders"


def test_default_template_validates_and_has_single_ip_table():
    tpl = LetterTemplate.model_validate(DEFAULT_TEMPLATE)
    ip_tables = [b for b in tpl.blocks if b.type == "ip_table"]
    assert len(ip_tables) == 1
    assert tpl.name == "IFSO Dwarka Default"


def test_template_rejects_two_ip_tables():
    import pytest
    from pydantic import ValidationError
    bad = {
        "name": "x", "scope": "user",
        "page": DEFAULT_TEMPLATE["page"],
        "blocks": [
            {"id": "a", "type": "ip_table"},
            {"id": "b", "type": "ip_table"},
        ],
    }
    with pytest.raises(ValidationError):
        LetterTemplate.model_validate(bad)


import pandas as pd
from docx import Document
from utils.isp_letter_generator import ISPLetterGenerator


def _sample_df():
    return pd.DataFrame([
        {"Type": "IPV4", "Search Value": "1.2.3.4",
         "From Date": "2025-01-28", "From Time": "14:30:25",
         "To Date": "2025-01-28", "To Time": "15:30:25"},
    ])


def test_build_ip_table_airtel_has_4_columns():
    doc = Document()
    ISPLetterGenerator()._build_ip_table(doc, "Airtel", _sample_df())
    table = doc.tables[0]
    assert len(table.columns) == 4
    headers = [table.rows[0].cells[i].text for i in range(4)]
    assert headers[0] == "Type" and headers[1] == "Search Value"


def test_build_ip_table_jio_has_6_columns_and_padded_time():
    doc = Document()
    ISPLetterGenerator()._build_ip_table(doc, "Jio", _sample_df())
    table = doc.tables[0]
    assert len(table.columns) == 6
    # From Time cell padded to 6 digits
    assert table.rows[1].cells[3].text == "143025"


def test_build_ip_table_vi_has_6_columns():
    doc = Document()
    ISPLetterGenerator()._build_ip_table(doc, "Vodafone Idea", _sample_df())
    assert len(doc.tables[0].columns) == 6


from utils.letter_template import DEFAULT_TEMPLATE


def _doc_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_render_substitutes_placeholders_in_text_and_list():
    case = {
        "fir_number": "201/25", "sections": "420 IPC", "fir_date": "01/02/2025",
        "police_station": "Special Cell", "complainant": "John Doe",
        "subject": "Reg info", "email_reference": "ref-1",
        "officer_name": "Insp X", "officer_designation": "IO",
        "officer_location": "Dwarka", "officer_contact": "999",
        "letter_date": "09/06/2026",
    }
    doc = ISPLetterGenerator().render_template_to_docx(
        DEFAULT_TEMPLATE, "Airtel", _sample_df(), case)
    text = _doc_text(doc)
    assert "FIR No.201/25" in text
    assert "Insp X" in text
    assert "{isp_name}" not in text          # all tokens replaced
    assert "Airtel" in text                  # isp_name injected
    # numbered list rendered
    assert "1. Details of the user" in text
    # the smart table is present with 4 cols (Airtel)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 4


def test_generate_letter_defaults_to_system_template_when_none_given():
    case = {"fir_number": "5/25", "subject": "Reg info", "email_reference": "",
            "police_station": "PS", "sections": "420", "fir_date": "1/1/25",
            "complainant": "X", "officer_name": "Y", "officer_designation": "IO",
            "officer_location": "L", "officer_contact": "9", "letter_date": "1/1/26"}
    doc = ISPLetterGenerator().generate_letter("Airtel", _sample_df(), case)
    text = _doc_text(doc)
    assert "Notice u/s 94 BNSS, 2023" in text     # default content present
    assert "FIR No.5/25" in text
    assert len(doc.tables[0].columns) == 4         # Airtel table preserved
