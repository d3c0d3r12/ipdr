import io

import mongomock
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.isp_letter_generator import ISPLetterGenerator
from utils.docx_import import parse_docx_to_blocks
from utils.letter_template import LetterTemplate
import pandas as pd


def _sample_df():
    return pd.DataFrame([
        {"Type": "IPV4", "Search Value": "1.2.3.4",
         "From Date": "2025-01-28", "From Time": "14:30:25",
         "To Date": "2025-01-28", "To Time": "15:30:25"},
    ])


def _docx_bytes(build) -> bytes:
    doc = Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _doc_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


# ---------- render_docx_template ----------

def _case():
    return {"fir_number": "201/25", "subject": "Reg info", "email_reference": "",
            "police_station": "PS", "sections": "420", "fir_date": "1/1/25",
            "complainant": "X", "officer_name": "Insp Y", "officer_designation": "IO",
            "officer_location": "L", "officer_contact": "9", "letter_date": "1/1/26"}


def test_render_docx_substitutes_tokens_and_injects_table_at_marker_airtel():
    def build(doc):
        doc.add_paragraph("Subject: FIR {fir_number} by {officer_name}")
        doc.add_paragraph("{ip_table}")
        doc.add_paragraph("Yours faithfully")
    out = ISPLetterGenerator().render_docx_template(_docx_bytes(build), "Airtel", _sample_df(), _case())
    text = _doc_text(out)
    assert "FIR 201/25 by Insp Y" in text
    assert "{ip_table}" not in text          # marker consumed
    assert "{fir_number}" not in text
    assert len(out.tables) == 1
    assert len(out.tables[0].columns) == 4   # Airtel 4-col


def test_render_docx_table_columns_for_jio():
    def build(doc):
        doc.add_paragraph("{ip_table}")
    out = ISPLetterGenerator().render_docx_template(_docx_bytes(build), "Jio", _sample_df(), _case())
    assert len(out.tables[0].columns) == 6   # Jio 6-col


def test_render_docx_appends_table_when_no_marker():
    def build(doc):
        doc.add_paragraph("A letter with no marker {fir_number}")
    out = ISPLetterGenerator().render_docx_template(_docx_bytes(build), "Airtel", _sample_df(), _case())
    assert len(out.tables) == 1
    assert "201/25" in _doc_text(out)


def test_generate_letter_routes_docx_kind():
    import base64
    def build(doc):
        doc.add_paragraph("Hello {officer_name}")
        doc.add_paragraph("{ip_table}")
    tpl = {"kind": "docx", "docx_b64": base64.b64encode(_docx_bytes(build)).decode(), "blocks": []}
    out = ISPLetterGenerator().generate_letter("Airtel", _sample_df(), _case(), tpl)
    assert "Hello Insp Y" in _doc_text(out)
    assert len(out.tables[0].columns) == 4


# ---------- parse_docx_to_blocks ----------

def test_parse_docx_to_blocks_text_list_table():
    def build(doc):
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run("OFFICE HEADER")
        r.bold = True
        doc.add_paragraph("First point", style="List Number")
        doc.add_paragraph("Second point", style="List Number")
        doc.add_table(rows=1, cols=4)
    blocks = parse_docx_to_blocks(_docx_bytes(build))
    types = [b["type"] for b in blocks]
    assert "text" in types and "list" in types and "ip_table" in types
    heading = next(b for b in blocks if b["type"] == "text")
    assert heading["align"] == "center" and heading["bold"] is True
    lst = next(b for b in blocks if b["type"] == "list")
    assert lst["items"] == ["First point", "Second point"]
    # result must validate as a blocks template
    LetterTemplate.model_validate({"name": "x", "kind": "blocks", "blocks": blocks})


def test_parse_docx_marker_becomes_ip_table_block():
    def build(doc):
        doc.add_paragraph("Subject {fir_number}")
        doc.add_paragraph("{ip_table}")
        doc.add_paragraph("Regards")
    blocks = parse_docx_to_blocks(_docx_bytes(build))
    types = [b["type"] for b in blocks]
    assert types == ["text", "ip_table", "text"]


def test_parse_docx_ignores_second_table():
    def build(doc):
        doc.add_paragraph("Hi")
        doc.add_table(rows=1, cols=4)
        doc.add_table(rows=1, cols=6)
    blocks = parse_docx_to_blocks(_docx_bytes(build))
    assert sum(1 for b in blocks if b["type"] == "ip_table") == 1


# ---------- upload-docx endpoint ----------

@pytest.fixture()
def client_with_db():
    from fastapi.testclient import TestClient
    import main
    from routers.auth_secure import get_current_user
    from core.db import get_db
    from services.letter_template_service import seed_default_template
    db = mongomock.MongoClient().testdb
    seed_default_template(db)
    main.app.dependency_overrides[get_current_user] = lambda: {"_id": "u1", "role": "investigator"}
    main.app.dependency_overrides[get_db] = lambda: db
    client = TestClient(main.app)  # no `with`: skip startup (avoid real Mongo)
    yield client, db
    main.app.dependency_overrides.clear()


def test_upload_docx_raw_creates_docx_template(client_with_db):
    client, db = client_with_db
    def build(doc):
        doc.add_paragraph("Letter {fir_number}")
    files = {"file": ("letter.docx", _docx_bytes(build),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = client.post("/api/letter-templates/upload-docx", files=files, data={"name": "My Word", "mode": "raw"})
    assert r.status_code == 200, r.text
    tpl = r.json()["template"]
    assert tpl["kind"] == "docx" and tpl["docx_b64"]


def test_upload_docx_convert_creates_blocks_template(client_with_db):
    client, db = client_with_db
    def build(doc):
        doc.add_paragraph("Heading")
        doc.add_table(rows=1, cols=4)
    files = {"file": ("letter.docx", _docx_bytes(build),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = client.post("/api/letter-templates/upload-docx", files=files, data={"name": "Converted", "mode": "convert"})
    assert r.status_code == 200, r.text
    tpl = r.json()["template"]
    assert tpl["kind"] == "blocks"
    assert any(b["type"] == "ip_table" for b in tpl["blocks"])


def test_upload_non_docx_rejected(client_with_db):
    client, _ = client_with_db
    files = {"file": ("notes.txt", b"hello", "text/plain")}
    r = client.post("/api/letter-templates/upload-docx", files=files, data={"name": "X", "mode": "raw"})
    assert r.status_code == 400


def test_upload_corrupt_docx_rejected(client_with_db):
    client, _ = client_with_db
    files = {"file": ("bad.docx", b"not a real docx", "application/octet-stream")}
    r = client.post("/api/letter-templates/upload-docx", files=files, data={"name": "X", "mode": "raw"})
    assert r.status_code == 422
